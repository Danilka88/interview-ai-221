import logging
import os
import tempfile
import io
import asyncio
from typing import Union
from pathlib import Path

from fastapi import UploadFile
import aiofiles

# --- Библиотеки для парсинга ---
import docx
import pypdf
from striprtf.striprtf import rtf_to_text

# --- Синхронные "воркеры" для CPU-bound операций ---

def _parse_pdf_sync(content: bytes) -> str:
    """Синхронно парсит содержимое PDF файла."""
    pdf_stream = io.BytesIO(content)
    reader = pypdf.PdfReader(pdf_stream)
    text = "".join(page.extract_text() + "\n" for page in reader.pages if page.extract_text())
    return text

def _parse_docx_sync(content: bytes) -> str:
    """Синхронно парсит содержимое DOCX файла."""
    doc_stream = io.BytesIO(content)
    doc = docx.Document(doc_stream)
    full_text = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    return "\n".join(full_text)

def _parse_rtf_sync(content: str) -> str:
    """Синхронно парсит содержимое RTF файла."""
    return rtf_to_text(content)


# --- Асинхронная логика парсинга файлов ---

async def _extract_text_from_pdf(file_path: Union[str, Path]) -> str:
    """Асинхронно читает PDF и выносит парсинг в отдельный поток."""
    try:
        async with aiofiles.open(file_path, 'rb') as f:
            content = await f.read()
        # Выполняем блокирующую операцию в отдельном потоке
        return await asyncio.to_thread(_parse_pdf_sync, content)
    except Exception as e:
        logging.error(f"Ошибка при обработке PDF {file_path}: {e}")
        return ""

async def _extract_text_from_docx(file_path: Union[str, Path]) -> str:
    """Асинхронно читает DOCX и выносит парсинг в отдельный поток."""
    try:
        async with aiofiles.open(file_path, 'rb') as f:
            content = await f.read()
        # Выполняем блокирующую операцию в отдельном потоке
        return await asyncio.to_thread(_parse_docx_sync, content)
    except Exception as e:
        logging.error(f"Ошибка при обработке DOCX {file_path}: {e}")
        return ""

async def _extract_text_from_rtf(file_path: Union[str, Path]) -> str:
    """Асинхронно читает RTF и выносит парсинг в отдельный поток."""
    try:
        async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = await f.read()
        # Выполняем блокирующую операцию в отдельном потоке
        return await asyncio.to_thread(_parse_rtf_sync, content)
    except Exception as e:
        logging.error(f"Ошибка при обработке RTF {file_path}: {e}")
        return ""

async def _extract_text_from_plain(file_path: Union[str, Path]) -> str:
    """Асинхронно извлекает текст из простого текстового файла."""
    try:
        async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return await f.read()
    except Exception as e:
        logging.error(f"Ошибка при асинхронном чтении текстового файла {file_path}: {e}")
        return ""

# Словарь диспетчеризации теперь содержит асинхронные функции
EXTRACTORS = {
    ".pdf": _extract_text_from_pdf,
    ".docx": _extract_text_from_docx,
    ".rtf": _extract_text_from_rtf,
    ".txt": _extract_text_from_plain,
    ".md": _extract_text_from_plain,
}

async def extract_text_from_file(file_path: Union[str, Path]) -> str:
    """Асинхронно выбирает нужный экстрактор и извлекает текст из файла."""
    path = Path(file_path)
    extractor = EXTRACTORS.get(path.suffix.lower())
    if extractor:
        logging.info(f"Используется асинхронный экстрактор для {path.suffix.lower()}...")
        # Вызываем асинхронную функцию через await
        return await extractor(path)
    logging.warning(f"Файл с расширением {path.suffix.lower()} не поддерживается.")
    return ""

async def save_upload_file_tmp(upload_file: UploadFile) -> str:
    """Асинхронно сохраняет загруженный файл во временный файл."""
    try:
        suffix = Path(upload_file.filename).suffix
        # Используем синхронный tempfile для получения уникального имени, но запись делаем асинхронно
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp_path = tmp.name
        
        async with aiofiles.open(tmp_path, 'wb') as out_file:
            content = await upload_file.read()
            await out_file.write(content)
        
        return tmp_path
    finally:
        await upload_file.close()

def cleanup_file(file_path: str):
    """Синхронно удаляет временный файл."""
    try:
        os.remove(file_path)
        logging.info(f"Временный файл {file_path} удален.")
    except OSError as e:
        logging.error(f"Ошибка при удалении файла {file_path}: {e}")